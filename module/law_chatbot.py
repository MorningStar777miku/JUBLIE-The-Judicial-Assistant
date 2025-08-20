import streamlit as st
import pandas as pd
import requests
from PyPDF2 import PdfReader
from groq import Groq
import os
import re
from textwrap import wrap
from dotenv import load_dotenv
from docx import Document
from fpdf import FPDF

load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)

def read_pdf(file):
    """Reads a PDF and extracts only law-relevant sections."""
    reader = PdfReader(file)
    relevant_text = []
    keywords = [
        "ipc section", "indian penal code", "constitution of india",
        "article", "fundamental right", "fundamental duty",
        "preamble", "schedule", "act", "legal", "law", "court", "justice"
    ]
    for page in reader.pages:
        text = page.extract_text() or ""
        text_lower = text.lower()
        if any(kw in text_lower for kw in keywords):
            relevant_text.append(text.strip())
    return "\n\n".join(relevant_text)

def is_law_related(text):
    """Checks if the text is related to IPC, Constitution, or other legal topics."""
    text_lower = text.lower()
    keywords = [
        "ipc section", "indian penal code", "constitution of india",
        "fundamental right", "fundamental duty", "article",
        "act", "legal", "law", "court", "justice"
    ]
    if re.search(r"\bsection \d+\b", text_lower):
        return True
    return any(keyword in text_lower for keyword in keywords)

def generate_summary(long_text):
    """Summarizes text in chunks to avoid token limit issues."""
    try:
        chunks = wrap(long_text, 5000)
        summaries = []
        for chunk in chunks:
            response = client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                messages=[
                    {"role": "system", "content": "You are a legal expert. Summarize the key legal points."},
                    {"role": "user", "content": chunk}
                ],
                max_tokens=5000
            )
            summaries.append(response.choices[0].message.content)

        final_response = client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=[
                {"role": "system", "content": "You are a legal expert. Merge these partial summaries into one coherent summary."},
                {"role": "user", "content": "\n\n".join(summaries)}
            ],
            max_tokens=5000
        )
        return final_response.choices[0].message.content

    except Exception as e:
        if "rate_limit_exceeded" in str(e).lower():
            return "You've reached today's token limit for legal summarization. Please try again later."
        return f"Error generating summary: {e}"

@st.cache_data
def load_coi_dataset():
    """Loads coi dataset."""
    project_root = os.path.dirname(os.path.dirname(__file__))
    csv_path = os.path.join(project_root, "dataset/coi.csv")
    if not os.path.exists(csv_path):
        raise FileNotFoundError(f"Could not find coi.csv at {csv_path}")
    return pd.read_csv(csv_path)


def export_chat(format_type="pdf"):
    messages = st.session_state.get("law_messages", [])
    if not messages:
        st.warning("No conversation to export.")
        return

    if format_type == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for msg in messages:
            pdf.multi_cell(0, 10, f"{msg['role'].capitalize()}: {msg['content']}")
        pdf.output("law_chat.pdf")
        with open("law_chat.pdf", "rb") as f:
            st.download_button("Download Chat as PDF", f, file_name="law_chat.pdf")

    elif format_type == "docx":
        doc = Document()
        doc.add_heading("Law Chat Conversation", level=1)
        for msg in messages:
            doc.add_paragraph(f"{msg['role'].capitalize()}: {msg['content']}")
        doc.save("law_chat.docx")
        with open("law_chat.docx", "rb") as f:
            st.download_button("Download Chat as DOCX", f, file_name="law_chat.docx")

def show_law_chatbot():
    st.subheader("Law jublie")
    coi_df = load_coi_dataset()

    with st.expander("Export Chat History"):
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Export as PDF"):
                export_chat("pdf")
        with col2:
            if st.button("Export as DOCX"):
                export_chat("docx")
    st.divider()

    if "law_messages" not in st.session_state:
        st.session_state.law_messages = []

    if "show_complaint_form" not in st.session_state:
        st.session_state.show_complaint_form = False

    for msg in st.session_state.law_messages:
        st.chat_message(msg["role"]).markdown(msg["content"])

    pdf_file = st.file_uploader("Upload a law-related PDF", type=["pdf"])
    if pdf_file:
        with st.spinner("Reading your PDF..."):
            file_text = read_pdf(pdf_file)
            if file_text.strip():
                if is_law_related(file_text):
                    summary = generate_summary(file_text)
                    st.session_state.law_messages.append({"role": "assistant", "content": summary})
                    st.chat_message("assistant").markdown(summary)
                else:
                    msg = "This PDF does not appear to be related to IPC or the Constitution of India."
                    st.session_state.law_messages.append({"role": "assistant", "content": msg})
                    st.chat_message("assistant").markdown(msg)
    st.divider()

    col1, col2 = st.columns([8, 2]) 
    with col1:
        law_query = st.chat_input("Type your law question here...", key="law_input")
    with col2:
        if st.button("Complaint"):
            st.session_state.show_complaint_form = not st.session_state.show_complaint_form
    if st.session_state.show_complaint_form:
        st.subheader("Raise a Complaint")
        with st.form("complaint_form"):
            complaint_title = st.text_input("Complaint Title")
            complaint_description = st.text_area("Complaint Description")
            submitted = st.form_submit_button("Submit Complaint")

            if submitted:
                if complaint_title and complaint_description:
                    st.success("Your complaint has been submitted successfully!")
                    with open("complaints.txt", "a") as f:
                        f.write(f"Title: {complaint_title}\nDescription: {complaint_description}\n\n")
                    st.session_state.show_complaint_form = False  
                else:
                    st.error("Please fill in both the title and description.")
    if law_query:
        st.session_state.law_messages.append({"role": "user", "content": law_query})
        with st.chat_message("user"):
            st.markdown(law_query)

        if not is_law_related(law_query):
            warning_msg = "This chatbot only handles Section and Constitution-related queries. Please use the General Chatbot for other questions."
            st.session_state.law_messages.append({"role": "assistant", "content": warning_msg})
            with st.chat_message("assistant"):
                st.markdown(warning_msg)
            return

        with st.spinner("Jublie is Thinking..."):
            headers = {
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": "meta-llama/llama-4-scout-17b-16e-instruct",
                "messages": [{"role": "user", "content": law_query}]
            }
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers=headers,
                json=payload
            ).json()

            if "choices" in response:
                generated_text = response["choices"][0]["message"]["content"]

                pattern = re.escape(generated_text.lower())
                matched = coi_df[
                    coi_df['article'].str.lower().str.contains(pattern, na=False, regex=True) |
                    coi_df['title'].str.lower().str.contains(pattern, na=False, regex=True) |
                    coi_df['description'].str.lower().str.contains(pattern, na=False, regex=True)
                ]
                if not matched.empty:
                    generated_text += f"\n\n**Verified Legal Content:**\n{matched.iloc[0]['description']}"

                st.session_state.law_messages.append({"role": "assistant", "content": generated_text})
                with st.chat_message("assistant"):
                    st.markdown(generated_text)
            else:
                error_msg = f"GROQ API Error: {response}"
                st.session_state.law_messages.append({"role": "assistant", "content": error_msg})
                with st.chat_message("assistant"):
                    st.markdown(error_msg)

